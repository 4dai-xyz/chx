#!/usr/bin/env python3
"""生成 SLAM 算法工程师面试准备 Word 文档。"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


REPO = Path("/home/ros/ros2_orbslam3")
OUT = REPO / "SLAM算法工程师面试准备_视觉SLAM_高斯重建_定位.docx"
JOB_IMAGE = REPO / "招聘要求.jpg"


def set_font(run, size: float = 10.5, bold: bool = False) -> None:
    """设置中英文字体。"""
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.bold = bold


def border(paragraph, color: str = "4F81BD") -> None:
    """给标题加底线。"""
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


def h1(doc: Document, text: str) -> None:
    """一级标题。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, 15, True)
    border(p)


def h2(doc: Document, text: str) -> None:
    """二级标题。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_font(r, 12.5, True)


def para(doc: Document, text: str, size: float = 10.0, bold: bool = False) -> None:
    """普通段落。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(text)
    set_font(r, size, bold)


def bullet(doc: Document, text: str, level: int = 0) -> None:
    """项目符号段落。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.left_indent = Cm(0.45 + 0.35 * level)
    p.paragraph_format.first_line_indent = Cm(-0.22)
    r = p.add_run("• " + text)
    set_font(r, 9.6)


def numbered(doc: Document, index: int, text: str) -> None:
    """编号段落。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.left_indent = Cm(0.55)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    r = p.add_run(f"{index}. {text}")
    set_font(r, 9.6)


def qa(doc: Document, question: str, answer: str) -> None:
    """问答块。"""
    p_q = doc.add_paragraph()
    p_q.paragraph_format.space_before = Pt(3)
    p_q.paragraph_format.space_after = Pt(1)
    r_q = p_q.add_run("Q: " + question)
    set_font(r_q, 10.0, True)
    p_a = doc.add_paragraph()
    p_a.paragraph_format.space_after = Pt(2)
    p_a.paragraph_format.left_indent = Cm(0.35)
    r_a = p_a.add_run("A: " + answer)
    set_font(r_a, 9.4)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    """添加简单表格。"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        run = table.rows[0].cells[i].paragraphs[0].add_run(header)
        set_font(run, 9.0, True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            run = cells[i].paragraphs[0].add_run(value)
            set_font(run, 8.6)
    doc.add_paragraph()


def build() -> None:
    """生成文档。"""
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.3)
    section.bottom_margin = Cm(1.3)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)

    for style_name in ("Normal", "Body Text", "List Paragraph"):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SLAM 算法工程师面试准备\n视觉 SLAM / 高斯重建 / 多传感器定位")
    set_font(run, 18, True)
    para(doc, "适用岗位：SLAM 算法工程师，重点覆盖视觉/视觉惯性定位、在线低延迟位姿跟踪、离线轨迹优化、稠密/语义建图、多视角协同定位、嵌入式部署与性能优化。", 10)

    if JOB_IMAGE.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(JOB_IMAGE), width=Cm(5.4))
        para(doc, "上图岗位要求摘要：多模态融合 SLAM、头戴/腕带视角高精度实时定位、头部位姿估计与手部/腕部轨迹重建、在线低延迟位姿跟踪、离线轨迹优化与语义建图、多视角时间同步/外参标定/坐标系统一，以及 Orin/RK3588 等嵌入式平台部署。", 8.8)

    h1(doc, "1. 岗位要求拆解与复习优先级")
    add_table(
        doc,
        ["岗位要求", "你要准备的知识", "结合你项目的说法"],
        [
            ["多模态融合 SLAM", "视觉几何、IMU 预积分、EKF/ESKF、因子图优化、时间同步、外参标定", "我已经做过 ORB-SLAM3、DPVO、VGGT 和 RGB-D/IMU 融合定位，理解单视觉和融合定位各自的优势与短板。"],
            ["头戴/腕带实时定位", "相机模型、位姿估计、低延迟 tracking、重定位、坐标系转换", "我的商场第一人称视频项目接近头戴视角，重点处理动态人群、反光、重复纹理和转弯失效。"],
            ["头部位姿与腕部轨迹重建", "SE(3) 位姿、相对位姿、手眼标定、多相机/多刚体坐标变换", "可以把头部相机作为世界定位源，把腕部相机/IMU 作为相对轨迹源，通过外参和时间同步统一到同一世界系。"],
            ["在线 SLAM 低延迟", "前端跟踪、关键帧策略、局部地图、滑窗优化、GPU/CPU 分工", "DPVO 在线 tracking 依赖深度网络和 CUDA，ORB-SLAM3 前端更轻；实际部署要根据硬件做降采样、关键帧稀疏化和异步线程。"],
            ["离线轨迹优化与语义建图", "BA、Pose Graph、Loop Closure、COLMAP、VGGT、3DGS、语义分割", "我做过 VGGT + COLMAP/BA + 3DGS 路线探索，理解在线定位和离线高质量建图的分工。"],
            ["多视角协同定位", "多传感器时间同步、外参标定、坐标系统一、协同 BA", "关键点是所有传感器要统一时间戳和坐标系，外参误差会直接变成定位系统误差。"],
            ["嵌入式部署", "TensorRT、ONNX、ROS2、内存/功耗/帧率、NPU/GPU", "我在 WSL/CUDA 环境跑通算法，后续到 Orin/RK3588 要做模型裁剪、推理加速和模块异步化。"],
        ],
    )

    h1(doc, "2. 面试开场：1 分钟项目自我介绍模板")
    para(doc, "建议回答：我最近主要做商场室内场景的视觉 SLAM 和具身导航数据采集。场景是第一人称视频，类似头戴相机，会遇到行人遮挡、反光地面、重复纹理、转弯视角变化和运动模糊。我搭建了 ROS2 + OpenCV + CUDA 的实验环境，跑通了 ORB-SLAM3、DPVO、Mono-VO/光流法，并测试了 VGGT 做离线三维重建。")
    para(doc, "在 DPVO 上，我做了后处理增强：读取 DPVO 输出轨迹，建立关键帧记忆，计算模糊度、遮挡比例、位姿变化等质量指标，针对 weak 片段用 ORB 特征匹配 + Essential Matrix + RANSAC 做几何验证和重定位候选检索。最终在商场遮挡视频上输出 1590 个位姿样本、158 个关键帧、172 个质量事件，其中几何验证通过比例约 95.3%。")
    para(doc, "对于离线建图，我测试了 VGGT 对视频抽帧后的相机位姿、深度、点云输出，并探索 COLMAP + BA + 3DGS 后处理。我的理解是：在线 SLAM 更关注实时、鲁棒、低延迟；离线重建更关注全局一致性、轨迹精度和地图质量。")

    h1(doc, "3. 视觉 SLAM 总体框架")
    para(doc, "视觉 SLAM 的目标是：在未知环境中，仅依靠相机或相机+IMU等传感器，同时估计传感器自身位姿和环境地图。核心状态通常是相机位姿 T_wc 或 T_cw、地图点 X_w、速度、IMU bias，以及相机/IMU外参。")
    h2(doc, "3.1 标准流程")
    numbered(doc, 1, "图像输入与预处理：去畸变、灰度化/归一化、降采样、时间戳对齐。")
    numbered(doc, 2, "前端跟踪：提取特征或直接使用光度误差/神经网络估计相邻帧位姿。")
    numbered(doc, 3, "局部建图：选择关键帧，三角化地图点，维护局部地图。")
    numbered(doc, 4, "后端优化：Bundle Adjustment、Pose Graph、滑窗优化或因子图优化。")
    numbered(doc, 5, "回环检测与重定位：识别曾经到过的位置，消除累计漂移。")
    numbered(doc, 6, "地图输出：稀疏点云、稠密点云、语义地图、3DGS/NeRF类可渲染地图。")
    h2(doc, "3.2 特征法、直接法、学习法对比")
    add_table(
        doc,
        ["方法", "代表", "优点", "缺点", "面试回答要点"],
        [
            ["特征法", "ORB-SLAM3", "鲁棒、解释性强、可回环、工程成熟", "弱纹理/动态物体/运动模糊下容易失效", "适合传统 SLAM 体系，关键点是 ORB 特征、BoW、PnP、BA。"],
            ["直接法/光流法", "DSO、LK 光流、Mono-VO", "可利用更多像素信息，弱纹理中有时比稀疏特征更连续", "对光照变化、曝光、动态物体敏感", "核心是光度一致性和小运动假设。"],
            ["学习法 VO", "DPVO、DROID-SLAM", "特征/匹配更强，可在复杂纹理中提升鲁棒性", "依赖 GPU、训练域、可解释性弱，部署复杂", "适合复杂室内视频，但要处理尺度、漂移和动态干扰。"],
            ["多视图几何模型", "VGGT", "可直接预测相机、深度、点云，适合离线几何重建", "长视频需分块对齐，实时性弱", "更像离线重建/地图生成，不是传统实时 SLAM。"],
        ],
    )

    h1(doc, "4. 几何基础：面试高频公式与概念")
    h2(doc, "4.1 相机模型")
    bullet(doc, "针孔模型：s [u, v, 1]^T = K [R | t] [X, Y, Z, 1]^T。K 包含 fx、fy、cx、cy。")
    bullet(doc, "畸变：常见有径向畸变 k1/k2/k3、切向畸变 p1/p2。鱼眼/广角相机还会用 Kannala-Brandt 模型。")
    bullet(doc, "单目尺度不可观：单目 SLAM 只能恢复相对尺度，若没有 IMU、深度、已知尺寸或地图约束，无法得到真实米制尺度。")
    h2(doc, "4.2 对极几何")
    bullet(doc, "Fundamental Matrix F：用于未标定相机，满足 x2^T F x1 = 0。")
    bullet(doc, "Essential Matrix E：用于已知内参的归一化坐标，满足 x2^T E x1 = 0，E = [t]_x R。")
    bullet(doc, "E 分解可以得到 R 和 t 的方向，但单目无法得到 t 的真实尺度。")
    bullet(doc, "RANSAC 的作用是从含有外点的匹配中估计稳定模型，剔除动态物体、错误匹配和重复纹理造成的外点。")
    h2(doc, "4.3 PnP、三角化和 BA")
    bullet(doc, "PnP：已知 3D 地图点和当前帧 2D 像素匹配，求相机位姿。重定位和 tracking 常用 PnP + RANSAC。")
    bullet(doc, "三角化：已知两个或多个视角的像素观测和相机位姿，恢复 3D 点。视差太小、匹配错误或相机运动退化会导致三角化不稳定。")
    bullet(doc, "Bundle Adjustment：联合优化相机位姿和地图点，最小化重投影误差 Σ ||u_ij - π(T_i, X_j)||²。")
    bullet(doc, "Pose Graph：只优化关键帧位姿和相对约束，常用于回环后全局漂移校正，比完整 BA 更轻。")

    h1(doc, "5. ORB-SLAM3 重点")
    para(doc, "ORB-SLAM3 是典型特征法 SLAM。它的主线是 ORB 特征提取与匹配、初始化、Tracking、Local Mapping、Loop Closing、Atlas 多地图管理、BA 优化和 Pangolin 可视化。")
    h2(doc, "5.1 ORB-SLAM3 流程")
    numbered(doc, 1, "Tracking 线程：从当前图像提取 ORB 特征，和上一帧/局部地图/关键帧匹配，估计当前相机位姿。")
    numbered(doc, 2, "Local Mapping 线程：插入关键帧、三角化新 MapPoint、剔除质量差的点、做局部 BA。")
    numbered(doc, 3, "Loop Closing 线程：用 BoW 做地点识别，Sim3 校正尺度漂移，做 pose graph 优化。")
    numbered(doc, 4, "Viewer 线程：显示当前帧、关键帧、地图点和轨迹。")
    h2(doc, "5.2 面试容易追问")
    bullet(doc, "为什么 ORB 特征适合实时 SLAM：FAST 角点 + BRIEF 描述子，旋转不变，二进制描述子匹配快。")
    bullet(doc, "初始化为什么难：单目需要足够视差，纯旋转、低纹理、动态物体会导致初始化失败或尺度不稳定。")
    bullet(doc, "回环为什么能减小漂移：用历史关键帧建立全局约束，优化整张 pose graph。")
    bullet(doc, "在商场为什么容易失败：人群是动态物体，地面反光破坏光度/特征稳定性，重复纹理导致错误匹配，转弯视角变化导致重叠区域少。")

    h1(doc, "6. 视觉惯性 SLAM：VINS / OpenVINS / ROVIO / ESKF")
    h2(doc, "6.1 IMU 为什么重要")
    bullet(doc, "IMU 提供高频角速度和加速度，可以补足相机低帧率、运动模糊和短时遮挡问题。")
    bullet(doc, "视觉提供低频但不漂太快的位置/姿态约束，IMU 会随时间积分漂移，两者互补。")
    bullet(doc, "视觉+IMU 可以恢复单目尺度、提高快速运动下的稳定性。")
    h2(doc, "6.2 IMU 预积分")
    para(doc, "IMU 原始数据频率高，如果每次优化都重新积分会很慢。预积分是在两个关键帧之间先把 IMU 测量积分成 ΔR、Δv、Δp，并考虑 bias 的一阶修正。优化时把这段 IMU 约束作为因子加入滑窗。")
    bullet(doc, "状态量：位置 p、速度 v、姿态 R、陀螺仪 bias bg、加速度计 bias ba。")
    bullet(doc, "误差来源：噪声、bias 漂移、时间同步误差、外参误差、重力方向初始化误差。")
    h2(doc, "6.3 EKF / ESKF 和图优化的区别")
    add_table(
        doc,
        ["方法", "特点", "优点", "不足"],
        [
            ["EKF/ESKF", "递推滤波，只保留当前状态和协方差", "实时性好，适合嵌入式", "线性化点固定，历史状态难重优化"],
            ["滑窗优化/VINS", "保留最近一段关键帧和 IMU 因子", "精度高，可联合优化 bias/外参/位姿", "计算量大，需要边缘化"],
            ["Pose Graph/全局 BA", "离线或低频全局优化", "能消除漂移，地图一致性好", "不适合每帧实时运行"],
        ],
    )

    h1(doc, "7. DPVO：你项目中最该讲清楚的部分")
    para(doc, "DPVO 是一种深度学习视觉里程计方法，核心思想是用神经网络提取 patch 特征和相关性，再结合可微 Bundle Adjustment / 图优化估计相机运动。它比纯 ORB 特征法更依赖 GPU，但对复杂图像匹配有更强表达能力。")
    h2(doc, "7.1 你已经做过的工作")
    bullet(doc, "跑通 DPVO CUDA 扩展和 PyTorch CUDA 环境，使用商场第一人称视频进行离线定位。")
    bullet(doc, "用遮挡后视频 input_video.mp4_bev 减少行人和地面反光干扰。")
    bullet(doc, "读取 DPVO 输出 TUM 轨迹，做质量评分、关键帧记录、weak/lost 事件分析。")
    bullet(doc, "使用 ORB + Essential Matrix + RANSAC 做 2D-2D 几何验证，判断 weak 片段是否还能和历史关键帧建立稳定几何关系。")
    bullet(doc, "输出 1590 个位姿样本、158 个关键帧、172 个质量事件，几何验证通过比例 95.3%。")
    h2(doc, "7.2 面试时可以承认的边界")
    bullet(doc, "当前增强模块不是重新训练 DPVO，也不是完整 PnP 重定位。")
    bullet(doc, "当前几何验证是 2D-2D，下一步要把关键帧特征与 DPVO 三维点绑定，做 2D-3D PnP + 局部 BA。")
    bullet(doc, "单目 DPVO 仍有尺度不确定、漂移和动态物体干扰问题。若要上产品，最好融合 IMU/轮速/深度或语义动态物体剔除。")

    h1(doc, "8. VGGT 与离线三维重建")
    para(doc, "VGGT 是视觉几何 Transformer 模型，可从多张图像直接预测相机位姿、深度、点云等几何量。它适合做离线三维理解和地图生成，不等同于传统实时 SLAM。")
    h2(doc, "8.1 你项目中的 VGGT 结果")
    bullet(doc, "输入视频 1920x1080，约 29.4 FPS，总帧数 3181。")
    bullet(doc, "每 30 帧抽 1 帧，共 107 帧；window-size=8，window-stride=4，完成 26 个窗口。")
    bullet(doc, "输出 camera_centers、深度图、置信度图、窗口点云和对齐后的 full scene。")
    bullet(doc, "长视频需要分块对齐，否则每个窗口的局部坐标系不同，会出现轨迹断裂或地图不连续。")
    h2(doc, "8.2 VGGT 和传统 SfM/COLMAP 的关系")
    bullet(doc, "COLMAP 是经典 SfM/MVS：特征匹配 -> 几何验证 -> 增量式重建 -> BA -> 稠密重建。")
    bullet(doc, "VGGT 更像学习式几何先验，可以快速预测位姿/深度，但长序列全局一致性仍需要 Sim(3) 对齐、BA 或 COLMAP 后处理。")
    bullet(doc, "面试可说：我把 VGGT 看成离线建图的几何初始化工具，后面可以接 COLMAP/BA 提高一致性，再接 3DGS 生成可渲染地图。")

    h1(doc, "9. 3D Gaussian Splatting / 高斯重建")
    para(doc, "3D Gaussian Splatting（3DGS）是一种可微渲染表示。场景由大量三维高斯组成，每个高斯有位置、尺度、旋转、不透明度和颜色/球谐系数。训练时通过可微渲染把高斯投影到图像上，最小化渲染图和真实图像之间的损失。")
    h2(doc, "9.1 3DGS 的核心参数")
    bullet(doc, "均值 μ：高斯中心位置。")
    bullet(doc, "协方差 Σ：控制高斯椭球的形状和方向，通常由 scale + rotation 参数化。")
    bullet(doc, "opacity α：不透明度，影响体渲染式 alpha blending。")
    bullet(doc, "color / SH：颜色或球谐系数，用于表达视角相关外观。")
    h2(doc, "9.2 3DGS 输入与训练流程")
    numbered(doc, 1, "输入：多视角图像、相机内参、相机外参/位姿、初始点云。")
    numbered(doc, 2, "初始化：通常来自 COLMAP 稀疏点云，也可以来自 VGGT/深度估计点云。")
    numbered(doc, 3, "渲染：将 3D 高斯投影到当前视角，进行 splatting 和 alpha 混合。")
    numbered(doc, 4, "损失：常用 L1 + SSIM，优化渲染图与真实图像的差异。")
    numbered(doc, 5, "densification/pruning：对高误差区域增加高斯，对低贡献高斯剪枝。")
    h2(doc, "9.3 面试常问：3DGS 和 SLAM 什么关系")
    bullet(doc, "SLAM 主要解决实时定位和地图一致性；3DGS 主要解决高质量新视角渲染和稠密外观建图。")
    bullet(doc, "3DGS 需要较好的相机位姿。如果位姿漂移很大，渲染会模糊、重影或出现毛刺。")
    bullet(doc, "在线系统可用 SLAM 估计实时位姿，离线用 BA 优化位姿，再训练 3DGS 生成高质量地图。")
    bullet(doc, "3DGS 不天然提供拓扑导航、语义理解或路径规划，需要额外加语义分割、占据栅格或导航图。")

    h1(doc, "10. 定位、重定位与多传感器融合")
    h2(doc, "10.1 Tracking、Localization、Relocalization 的区别")
    bullet(doc, "Tracking：连续帧之间估计运动，依赖上一帧/局部地图。")
    bullet(doc, "Localization：在已有地图中估计当前位姿，可以是连续的，也可以从任意位置开始。")
    bullet(doc, "Relocalization：tracking 丢失后，通过地点识别和 PnP 找回位姿。")
    h2(doc, "10.2 重定位标准流程")
    numbered(doc, 1, "当前帧提取特征或全局描述子。")
    numbered(doc, 2, "从关键帧数据库检索候选位置。")
    numbered(doc, 3, "做特征匹配或 learned matching。")
    numbered(doc, 4, "如果有 3D 地图点，用 PnP + RANSAC 估计位姿；如果只有 2D-2D，可先用 E/F 矩阵验证几何一致性。")
    numbered(doc, 5, "位姿成功后，恢复局部 tracking，并可做局部 BA。")
    h2(doc, "10.3 多视角/多传感器关键难点")
    bullet(doc, "时间同步：不同相机/IMU 时间戳不一致会导致运动畸变和融合误差。硬同步最好，软同步要插值和延迟估计。")
    bullet(doc, "外参标定：相机-IMU、头部-腕部、相机-机器人 base_link 的外参误差会直接映射成定位误差。")
    bullet(doc, "坐标系统一：要清楚 world、camera、imu、body、map、odom、base_link 的变换方向。")
    bullet(doc, "多视角协同：多相机可以提升可观测性，但要处理视角重叠、外参、时间同步和联合优化。")

    h1(doc, "11. 在线系统和离线后处理的工程架构")
    h2(doc, "11.1 在线低延迟 SLAM")
    bullet(doc, "目标：每帧快速输出当前位姿，延迟可控，失败可恢复。")
    bullet(doc, "常用策略：图像降采样、特征数限制、关键帧稀疏化、局部地图、滑窗优化、异步线程、GPU 推理加速。")
    bullet(doc, "常见性能指标：FPS、端到端延迟、ATE/RPE、跟踪丢失次数、重定位耗时、CPU/GPU 占用、内存、功耗。")
    h2(doc, "11.2 离线全局优化和建图")
    bullet(doc, "目标：不追求每帧低延迟，追求全局一致性、地图质量和轨迹精度。")
    bullet(doc, "常用模块：全局 BA、Pose Graph、Loop Closure、COLMAP、MVS、VGGT 对齐、3DGS 训练、语义地图生成。")
    bullet(doc, "适合岗位第 4 条：离线轨迹优化与稠密/语义建图，对在线结果做全局优化与后处理。")

    h1(doc, "12. 嵌入式部署：Orin / RK3588 相关准备")
    bullet(doc, "Orin：NVIDIA GPU + TensorRT 生态较强，适合深度模型推理、视觉前端和部分 CUDA 加速。")
    bullet(doc, "RK3588：有 NPU，但部署 PyTorch 模型通常要走 ONNX/RKNN 转换，算子支持和量化误差需要验证。")
    bullet(doc, "部署优化思路：模型轻量化、FP16/INT8、TensorRT/ONNX、固定输入尺寸、减少 Python 开销、C++/ROS2 节点化、零拷贝图像传输、异步队列。")
    bullet(doc, "SLAM 系统要注意线程和缓存：图像输入、前端 tracking、后端优化、地图维护、可视化和日志最好异步，不要让可视化阻塞定位。")

    h1(doc, "13. 结合你项目的高质量回答模板")
    qa(doc, "你做的 DPVO 增强到底增强了什么？",
       "我没有改 DPVO 官方模型，而是在 DPVO 输出轨迹之后做了工程增强：读取轨迹和视频，建立关键帧记忆，计算模糊度、遮挡比例、位姿变化等质量指标，标记 weak/lost 片段；对弱跟踪片段提取 ORB 特征，与历史关键帧匹配，用 Essential Matrix + RANSAC 做 2D-2D 几何验证，判断是否还能建立稳定几何关系。下一步可以把关键帧特征与 3D 地图点绑定，做完整 PnP + 局部 BA。")
    qa(doc, "为什么商场场景 SLAM 难？",
       "商场有大量动态行人，地面反光会破坏光度和特征稳定性，货架/地砖/灯带有重复纹理，转弯时连续帧重叠区域突然减少，第一人称视频还有运动模糊。这些都会造成特征误匹配、跟踪质量下降、初始化不稳定和重定位困难。")
    qa(doc, "如果让你继续提升系统鲁棒性，你会怎么做？",
       "短期我会做三件事：第一，语义或颜色遮挡预处理，降低行人和反光区域对匹配的影响；第二，把当前 2D-2D 几何验证升级为 2D-3D PnP 重定位，并在局部关键帧上做 BA；第三，引入 IMU 或轮速做视觉惯性融合，解决快速转弯、短时遮挡和单目尺度问题。长期会做多传感器时间同步、外参标定、语义地图和离线全局优化。")
    qa(doc, "你怎么看 ORB-SLAM3 和 DPVO 的取舍？",
       "ORB-SLAM3 是成熟的传统特征法系统，回环、局部建图、BA 体系完整，解释性强，CPU 也能跑；但在动态、弱纹理、反光和运动模糊中容易掉。DPVO 通过深度网络学习匹配和运动估计，复杂场景的表达能力更强，但依赖 GPU、训练域和模型部署，回环/地图管理也需要额外工程补齐。工程上可以用 DPVO 提供更强前端，用传统几何做验证、重定位和全局优化。")
    qa(doc, "你为什么测试 VGGT？它和 SLAM 有什么关系？",
       "VGGT 可以从多视角图像直接预测相机位姿、深度和点云，适合做离线几何重建和地图初始化。它不是传统实时 SLAM，但可以给离线建图提供很好的几何先验。我的理解是在线系统用 SLAM 提供实时定位，离线用 VGGT/COLMAP/BA/3DGS 提升地图质量和可视化效果。")
    qa(doc, "你做的 3DGS 为什么效果会毛刺或断裂？",
       "3DGS 对相机位姿质量很敏感。如果输入轨迹存在漂移、窗口对齐不连续、尺度不一致或动态物体没有剔除，训练出来的高斯会出现重影、毛刺、墙面不完整。要改善需要更好的相机位姿，比如 COLMAP/BA 优化、剔除动态物体、选取清晰关键帧、统一尺度并增加多视角覆盖。")
    qa(doc, "单目 SLAM 为什么没有真实尺度？怎么解决？",
       "单目投影中，如果场景和相机平移同时按同一比例缩放，图像观测不变，所以尺度不可观。解决方法包括融合 IMU、使用 RGB-D/双目、引入已知物体尺寸、地面约束、机器人里程计或先验地图。")
    qa(doc, "Essential Matrix 和 PnP 分别解决什么问题？",
       "Essential Matrix 用两帧 2D-2D 匹配估计相对姿态方向和几何一致性，适合验证两帧是否符合刚体运动；PnP 用 3D 地图点和当前 2D 像素匹配估计当前相机在地图中的绝对位姿，是重定位和 map-based localization 的核心。")
    qa(doc, "为什么需要时间同步和外参标定？",
       "多传感器融合时，每个传感器的数据必须对应同一时刻和同一坐标系。时间不同步会把运动误差当成测量误差，外参不准会让相机、IMU、机器人 base 的位姿互相对不上，最后表现为轨迹抖动、尺度错误或系统性偏移。")
    qa(doc, "如果岗位要你做头部位姿和腕部轨迹重建，你怎么设计？",
       "我会把头部相机/IMU作为主要世界定位源，估计 T_world_head；腕部相机/IMU估计 T_head_wrist 或 T_world_wrist。关键是头-腕之间的外参/相对位姿约束、时间同步和统一坐标系。如果腕部有相机，可用视觉惯性里程计估计相对轨迹，并通过头部地图或共同观测做约束融合。")

    h1(doc, "14. 刚才面试暴露的重点问题补充")
    h2(doc, "14.1 商场室内定位如何初始化")
    para(doc, "商场室内定位初始化，本质上是回答两个问题：第一，系统刚启动时如何确定当前相机在地图里的初始位姿；第二，如果还没有先验地图，如何从视频前几秒建立一个稳定的局部坐标系。面试时不要只说“跑 SLAM 初始化”，要区分有地图和无地图两种情况。")
    h2(doc, "情况 A：没有先验地图，只做在线 SLAM/VO 初始化")
    numbered(doc, 1, "相机启动后先做图像质量检查：曝光、模糊度、特征数量、动态物体比例。商场里人多、反光强，如果开头几帧质量差，可以延迟初始化或提示用户慢速移动。")
    numbered(doc, 2, "选择特征丰富、视差足够的连续帧做初始化。单目系统需要平移视差，纯旋转、站着不动、对着白墙或玻璃反光区域都不适合初始化。")
    numbered(doc, 3, "用两帧/多帧几何估计初始相对位姿：特征法可用 Homography / Essential Matrix 模型选择，学习式 VO 如 DPVO 可以直接输出初始轨迹，但仍要检查几何一致性。")
    numbered(doc, 4, "三角化初始地图点或建立初始关键帧窗口，再进入 tracking。初始化后继续用局部 BA 或滑窗优化稳定早期轨迹。")
    numbered(doc, 5, "如果使用 IMU，则通过静止或短时间数据估计重力方向、陀螺仪/加速度计 bias，并用 VIO 恢复尺度和重力对齐。")
    h2(doc, "情况 B：有先验地图，要做地图内定位初始化")
    numbered(doc, 1, "先验地图可以来自离线建图：ORB-SLAM3/COLMAP/VGGT/3DGS 的关键帧、稀疏点云、语义地标、店铺招牌、入口、电梯等。")
    numbered(doc, 2, "当前帧提取局部特征或全局描述子，与地图关键帧库做地点识别，找到候选区域。")
    numbered(doc, 3, "当前帧 2D 特征与地图 3D 点匹配，用 PnP + RANSAC 求 T_world_camera。若 3D 点不足，可先用 2D-2D Essential Matrix 做几何验证，再请求更多帧。")
    numbered(doc, 4, "初始化成功后，用局部地图跟踪和局部 BA 继续优化；若置信度低，则保持“未初始化/低置信度”状态，避免把错误位姿传给导航系统。")
    numbered(doc, 5, "商场可以结合语义地标提升初始化：店铺 Logo/OCR、扶梯/电梯、走廊结构、入口、地标广告牌。语义能缩小候选区域，几何负责精确位姿。")
    h2(doc, "商场初始化的工程策略")
    bullet(doc, "启动引导：眼镜端提示用户缓慢向前走 1-2 秒或轻微左右扫视，避免纯旋转和过多人群遮挡。")
    bullet(doc, "质量门控：特征数、RANSAC 内点数、内点比例、视差、IMU 运动幅度、图像模糊度都要过阈值再宣布初始化成功。")
    bullet(doc, "多候选机制：商场重复纹理多，单个候选容易误匹配，应该保留 top-k 候选，用几何验证、语义验证和连续帧一致性筛选。")
    bullet(doc, "失败恢复：初始化失败不应崩溃，进入 relocalization 状态，继续采集帧并寻找更好的视角。")
    qa(doc, "商场室内定位初始化你会怎么做？",
       "我会先区分有没有先验地图。如果没有地图，就是标准在线 SLAM 初始化：先做图像质量检查和运动引导，选择特征丰富且有足够平移视差的帧，用 Essential Matrix/Homography 或学习式 VO 建初始位姿，再三角化初始地图点并做局部优化。如果有商场地图，我会走 map-based localization：当前帧先用全局描述子或 BoW 检索候选关键帧，再做 2D-3D PnP + RANSAC 求初始位姿，最后用连续帧一致性和局部 BA 稳定。商场里我会额外利用 OCR/店铺招牌/扶梯电梯等语义地标来缩小候选范围。")
    qa(doc, "商场初始化为什么容易失败？",
       "主要是动态人群、反光地面、重复纹理、玻璃、低纹理墙面和纯旋转运动。单目初始化需要足够视差，如果用户刚戴上眼镜只是原地转头，深度不可观；如果开头画面都是行人或反光地面，匹配会被污染。所以工程上要做质量门控和初始化引导，必要时等到更稳定的几帧再初始化。")

    h2(doc, "14.2 VIO 和 IMU：面试必会回答")
    para(doc, "VIO 是 Visual-Inertial Odometry，视觉惯性里程计。它融合相机和 IMU：相机提供低频但相对稳定的几何观测，IMU 提供高频角速度和加速度，帮助处理快速运动、短时遮挡和单目尺度不可观问题。")
    h2(doc, "IMU 数据包含什么")
    bullet(doc, "陀螺仪 gyroscope：测角速度 ω，用于积分姿态。")
    bullet(doc, "加速度计 accelerometer：测比力，包含运动加速度和重力影响，用于估计速度、位置和重力方向。")
    bullet(doc, "IMU 误差：白噪声、bias、bias random walk、温漂、安装误差、时间戳误差。")
    h2(doc, "VIO 系统一般怎么做")
    numbered(doc, 1, "时间同步：相机帧和 IMU 高频数据必须对齐。")
    numbered(doc, 2, "外参标定：求 T_camera_imu，即相机和 IMU 的固定刚体变换。")
    numbered(doc, 3, "初始化：估计重力方向、尺度、速度、陀螺仪 bias 和加速度计 bias。")
    numbered(doc, 4, "IMU 预积分：把两个关键帧之间的高频 IMU 积分成 ΔR、Δv、Δp，作为优化因子。")
    numbered(doc, 5, "视觉约束：特征重投影误差或直接法光度误差。")
    numbered(doc, 6, "后端优化：EKF/ESKF 或滑窗非线性优化，联合估计位姿、速度、bias 和地图点。")
    h2(doc, "VIO 和纯视觉相比的优势")
    bullet(doc, "恢复单目尺度：纯单目只有相对尺度，视觉惯性可以通过加速度和重力约束恢复米制尺度。")
    bullet(doc, "抗快速运动：眼镜端转头快、画面模糊时，IMU 还能提供短时姿态预测。")
    bullet(doc, "抗短时遮挡：视觉跟踪短暂失败时，IMU 可以做状态传播，等待视觉恢复。")
    bullet(doc, "重力对齐：可以把轨迹的 z 轴和真实重力方向对齐，避免三维轨迹看起来“斜着飞”。")
    h2(doc, "VIO 难点")
    bullet(doc, "时间同步：几毫秒误差在快速转头时会造成明显位姿误差。")
    bullet(doc, "外参标定：相机和 IMU 的旋转/平移外参不准，会让融合状态系统性偏移。")
    bullet(doc, "bias 估计：IMU bias 不估计会积分发散。")
    bullet(doc, "初始化退化：静止太久、纯旋转、低加速度激励会导致尺度和 bias 不好估。")
    bullet(doc, "滚动快门：眼镜端小相机常有 rolling shutter，快速运动下图像几何会变形。")
    qa(doc, "你有没有接触过 VIO 和 IMU？",
       "我之前在临舟科技的室内定位项目里做过视觉和 IMU 的 EKF 融合，理解 IMU 的角速度、加速度、bias 和漂移问题。在现在的 SLAM 项目里，我主要跑的是单目 ORB-SLAM3、DPVO 和 VGGT，但我清楚如果要做眼镜端实时定位，VIO 是很关键的一步：IMU 可以提供高频姿态预测、恢复单目尺度、短时补偿视觉遮挡。工程上要重点解决相机-IMU 时间同步、外参标定、IMU 预积分、重力/bias 初始化和滑窗优化。")
    qa(doc, "如果让你把 DPVO 接 IMU，你会怎么做？",
       "我会把 DPVO 当成视觉前端，输出相邻帧或关键帧的视觉位姿/匹配约束；IMU 作为高频预测和尺度/重力约束。短期工程可以先做松耦合：DPVO 输出位姿，IMU/EKF 做状态传播和平滑。更完整的方案是紧耦合：在滑窗里同时放视觉重投影/相对位姿因子和 IMU 预积分因子，联合优化 pose、velocity、bias。")

    h2(doc, "14.3 第一人称眼镜端数据如何在线处理")
    para(doc, "眼镜端第一人称数据的特点是：视角跟随头部快速转动、运动模糊明显、画面里动态人多、算力和功耗受限、实时性要求高。回答时要体现系统架构思维，而不是只说把视频丢给 SLAM。")
    h2(doc, "在线处理架构")
    numbered(doc, 1, "传感器采集层：RGB/RGB-D 相机、IMU、可能还有麦克风/蓝牙/WiFi/气压计。每个数据都要有统一时间戳。")
    numbered(doc, 2, "预处理层：图像去畸变、缩放、曝光/模糊检测、动态区域 mask、rolling shutter 或运动模糊处理。")
    numbered(doc, 3, "前端 tracking：轻量视觉前端或 VIO 前端，输出高频当前位姿。可用 ORB/光流/学习式 VO，根据硬件选择。")
    numbered(doc, 4, "局部地图层：维护最近关键帧、局部点云、语义地标，支持重定位和局部 BA。")
    numbered(doc, 5, "语义理解层：OCR、店铺招牌、可通行区域、动态人群、扶梯/电梯等，用于辅助定位和导航。")
    numbered(doc, 6, "后端优化层：在线只做小窗口优化，离线或低频线程做全局 BA、回环和地图更新。")
    numbered(doc, 7, "输出层：向导航/VLA/AR 显示输出当前位姿、置信度、局部路径、地图锚点和异常状态。")
    h2(doc, "眼镜端在线系统的关键工程点")
    bullet(doc, "低延迟优先：眼镜端导航要先保证 tracking 低延迟，复杂建图可以放到后台或云端。")
    bullet(doc, "多线程/异步：采集、tracking、建图、语义、可视化、日志分线程，避免语义大模型阻塞定位。")
    bullet(doc, "置信度输出：不要只输出 pose，还要输出定位质量，比如特征数、内点比例、重定位状态、IMU/视觉一致性。")
    bullet(doc, "动态物体处理：人群区域可以用语义分割或光流/RANSAC 外点降低权重。")
    bullet(doc, "地图策略：在线维护轻量关键帧地图；高质量稠密地图、3DGS 或语义地图可以离线生成。")
    bullet(doc, "端云协同：眼镜端跑实时定位和局部避障，云端/工作站做全局地图优化、VGGT/3DGS 重建和数据回流。")
    h2(doc, "如果结合招聘岗位，可以这么讲")
    para(doc, "我会把眼镜端系统拆成在线定位和离线建图两条线。在线部分用 VIO/轻量 SLAM 保证低延迟位姿输出，IMU 做高频预测，视觉做漂移校正和重定位；离线部分用关键帧、COLMAP/BA、VGGT/3DGS 和语义分割提升地图质量。商场场景中，OCR/店铺招牌/扶梯/电梯等语义地标可以帮助地点识别，动态人群和反光地面需要做 mask 或降低权重。最终给 VLA 或导航模块的不只是图像，而是结构化状态：当前位姿、置信度、局部地图、可通行区域、语义锚点和异常状态。")
    qa(doc, "如何在线处理第一人称眼镜端数据？",
       "我会分层处理：底层先保证 RGB/IMU 时间同步和外参标定；图像先做去畸变、缩放、模糊/曝光质量检查和动态区域处理；前端用轻量 VIO 或视觉 SLAM 做低延迟 tracking，IMU 提供高频姿态预测；局部地图维护关键帧和地图点，用 PnP/RANSAC 做重定位；语义线程异步跑 OCR、可通行区域、动态物体检测，辅助定位和导航；离线或后台再做 BA、回环、VGGT/3DGS 高质量建图。眼镜端最重要的是低延迟、置信度输出和失败恢复，不能让重模型阻塞实时定位。")
    qa(doc, "眼镜端和普通手持视频 SLAM 有什么不同？",
       "眼镜端更接近头部运动，转头速度快、rolling shutter 和运动模糊更明显；视角高度和姿态有规律，但人群遮挡很多；设备算力、功耗和散热更受限；同时它要服务 AR 导航，所以位姿延迟和稳定性比离线精度更重要。工程上需要 IMU 高频预测、轻量化 tracking、异步语义处理和稳定的置信度/失败恢复机制。")

    h1(doc, "15. 高频面试题清单")
    qas = [
        ("ORB 特征为什么具有旋转不变性？", "ORB 在 FAST 角点上计算方向，并旋转 BRIEF 采样模式得到 rBRIEF，因此对图像旋转更鲁棒。"),
        ("BA 的优化变量和误差项是什么？", "变量是相机位姿和地图点，误差项是观测像素与三维点投影像素之间的重投影误差。"),
        ("回环检测会引入什么风险？", "错误回环会把地图拉坏，所以需要 BoW 候选、几何验证和一致性检查。"),
        ("纯旋转为什么无法三角化？", "没有平移基线就没有视差，深度不可恢复。"),
        ("RANSAC 阈值怎么选？", "与像素噪声、特征精度、分辨率和内参有关。阈值太小会误删内点，太大会保留外点。"),
        ("光流法的基本假设是什么？", "亮度恒定、小运动、局部平滑。快速运动、光照变化和遮挡会破坏假设。"),
        ("IMU bias 为什么要估计？", "陀螺仪和加速度计 bias 会随时间积分成越来越大的姿态、速度和位置误差。"),
        ("滑窗边缘化是什么？", "为了控制计算量，把旧状态从优化窗口中移除，同时用先验因子保留其约束信息。"),
        ("地图点如何剔除？", "根据观测次数、重投影误差、视角一致性、深度范围和跟踪成功率剔除。"),
        ("动态物体怎么处理？", "可用语义分割、光流不一致、RANSAC 外点、深度/运动一致性检测，把动态区域降低权重或剔除。"),
        ("多相机外参怎么标定？", "用标定板或自然特征，通过共同观测估计相机间固定变换，也可在 SLAM 中联合优化外参。"),
        ("ATE 和 RPE 区别？", "ATE 衡量全局轨迹误差，RPE 衡量局部相对运动误差。ATE 看全局漂移，RPE 看短时跟踪质量。"),
        ("3DGS 需要哪些输入？", "多视角图像、相机内参、相机位姿、初始点云。位姿越准，训练越稳定。"),
        ("为什么 3DGS 不是直接可导航地图？", "它是可渲染外观表示，不天然包含可通行区域、拓扑关系、语义和碰撞信息。需要转占据/语义/拓扑地图。"),
        ("你项目中 weak/lost 是怎么理解的？", "不是 DPVO 官方状态，而是我根据质量分数、遮挡、模糊、运动突变等指标做的后处理标签，用于分析跟踪风险和重定位候选。"),
        ("如果部署在 Orin，你会怎么优化 DPVO/VGGT？", "优先固定输入尺寸、FP16/TensorRT、减少窗口/关键帧数量、异步线程、降低可视化开销，必要时只把重模型放离线建图，在线用轻量 tracking。"),
    ]
    for q, a in qas:
        qa(doc, q, a)

    h1(doc, "16. 面试最后反问建议")
    bullet(doc, "目前系统主要面向头戴相机、腕带相机，还是多设备协同？是否已经有硬件同步方案？")
    bullet(doc, "在线定位对延迟和频率的目标是多少，例如 30Hz、60Hz 或端到端 <30ms？")
    bullet(doc, "公司当前更需要传统几何 SLAM、学习式 SLAM，还是离线高质量建图/数据生产管线？")
    bullet(doc, "Orin/RK3588 上是否已有推理框架和 ROS2 通信框架？视觉模型是否需要 TensorRT/RKNN 部署？")
    bullet(doc, "语义建图更偏目标级语义、可通行区域、人体/手部轨迹，还是用于大模型/VLA 训练数据标注？")

    h1(doc, "17. 面试前 10 分钟速记")
    bullet(doc, "SLAM = 前端跟踪 + 后端优化 + 回环/重定位 + 地图维护。")
    bullet(doc, "ORB-SLAM3：特征法，BoW 重定位/回环，局部 BA，全局 pose graph。")
    bullet(doc, "VINS/OpenVINS：视觉 + IMU，IMU 预积分/ESKF/滑窗，解决尺度和快速运动。")
    bullet(doc, "DPVO：学习式视觉里程计，GPU 依赖强，复杂图像匹配能力强；我做了关键帧、质量评分、几何验证。")
    bullet(doc, "VGGT：离线多视图几何模型，输出相机、深度、点云；适合建图初始化和三维理解。")
    bullet(doc, "3DGS：输入多视角图像和相机位姿，优化 3D 高斯做可微渲染；位姿不准会毛刺/重影。")
    bullet(doc, "定位鲁棒性关键：时间同步、外参标定、动态物体剔除、重定位、IMU 融合、局部/全局优化。")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
