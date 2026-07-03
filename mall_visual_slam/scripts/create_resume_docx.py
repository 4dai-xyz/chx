#!/usr/bin/env python3
"""根据 Markdown 简历生成可投递的 Word 文档。"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


REPO = Path("/home/ros/ros2_orbslam3")
SOURCE = REPO / "程浩轩-北京航空航天大学-机械硕士-2026暑期实习-SLAM机器人方向更新版.md"
TARGET = REPO / "程浩轩-北京航空航天大学-机械硕士-2026暑期实习-SLAM机器人方向更新版.docx"


def clean_inline(text: str) -> str:
    """去掉 Markdown 行内格式，保留纯文本。"""
    text = text.replace("**", "")
    text = text.replace("*", "")
    text = text.replace("`", "")
    return text.strip()


def set_font(run, size: float, bold: bool = False) -> None:
    """设置中英文字体。"""
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.bold = bold


def add_bottom_border(paragraph) -> None:
    """给章节标题添加底部细线。"""
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "4F81BD")
    borders.append(bottom)


def add_paragraph(doc: Document, text: str, size: float = 9.0, bold: bool = False, indent: bool = False):
    """添加普通段落。"""
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(1.5)
    paragraph.paragraph_format.line_spacing = 1.0
    if indent:
        paragraph.paragraph_format.left_indent = Cm(0.35)
        paragraph.paragraph_format.first_line_indent = Cm(-0.18)
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold)
    return paragraph


def add_bullet(doc: Document, text: str) -> None:
    """添加简历项目符号。"""
    paragraph = add_paragraph(doc, "• " + clean_inline(text), size=8.4, indent=True)
    paragraph.paragraph_format.space_after = Pt(1)


def add_section(doc: Document, title: str) -> None:
    """添加一级章节标题。"""
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(title)
    set_font(run, size=11.5, bold=True)
    add_bottom_border(paragraph)


def add_project_heading(doc: Document, text: str) -> None:
    """添加项目/经历标题。"""
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(1)
    run = paragraph.add_run(clean_inline(text))
    set_font(run, size=9.2, bold=True)


def build_docx() -> None:
    """生成 Word 简历。"""
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.05)
    section.bottom_margin = Cm(1.05)
    section.left_margin = Cm(1.15)
    section.right_margin = Cm(1.15)

    for style_name in ["Normal", "Body Text", "List Paragraph"]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(9)

    for raw_line in SOURCE.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("# "):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(1)
            run = paragraph.add_run(clean_inline(line[2:]))
            set_font(run, size=17, bold=True)
        elif line.startswith("**求职意向"):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(2)
            run = paragraph.add_run(clean_inline(line))
            set_font(run, size=9.2, bold=False)
        elif line.startswith("## "):
            add_section(doc, clean_inline(line[3:]))
        elif line.startswith("### "):
            add_project_heading(doc, line[4:])
        elif line.startswith("- "):
            add_bullet(doc, line[2:])
        elif re.match(r"^\*\*.+\*\*.*", line):
            add_project_heading(doc, line)
        else:
            add_paragraph(doc, clean_inline(line), size=8.8)

    doc.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    build_docx()
