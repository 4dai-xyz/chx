#!/usr/bin/env python3
"""生成保留原简历结构的 Word 版本。

这个脚本不使用前一版重新包装的 SLAM 简历模板，而是按原 PDF 的结构生成：
教育经历 -> 竞赛经历 -> 实习项目经验 -> 专业技能 -> 证书奖项及学术成果。
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


REPO = Path("/home/ros/ros2_orbslam3")
SOURCE = REPO / "程浩轩-北京航空航天大学-机械硕士-2026暑期实习-原格式保留项目更新版.md"
PHOTO = REPO / "简历原图_X5.jpg"
TARGET = REPO / "程浩轩-北京航空航天大学-机械硕士-2026暑期实习-原格式保留项目更新版.docx"


def set_run_font(run, size: float, bold: bool = False) -> None:
    """统一设置中英文字体。"""
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.bold = bold


def set_cell_border(cell, color: str = "FFFFFF") -> None:
    """隐藏表格边框。"""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "0")
        element.set(qn("w:color"), color)


def add_section_heading(doc: Document, text: str) -> None:
    """添加类似原简历的章节标题。"""
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(1.5)
    run = paragraph.add_run(text)
    set_run_font(run, 10.5, bold=True)
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "A6A6A6")
    borders.append(bottom)
    p_pr.append(borders)


def add_plain_paragraph(doc: Document, text: str, size: float = 7.6, bold: bool = False) -> None:
    """添加普通内容段落。"""
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0.8)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    set_run_font(run, size, bold=bold)


def parse_markdown_lines() -> list[str]:
    """读取并做少量 Markdown 标记清理。"""
    cleaned: list[str] = []
    for raw_line in SOURCE.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        line = line.replace("**", "")
        cleaned.append(line)
    return cleaned


def add_header(doc: Document, lines: list[str]) -> int:
    """添加姓名、联系方式和原头像。返回正文起始行号。"""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(13.0)
    table.columns[1].width = Cm(2.6)
    left = table.cell(0, 0)
    right = table.cell(0, 1)
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_border(left)
    set_cell_border(right)

    name = lines[0].lstrip("#").strip()
    p_name = left.paragraphs[0]
    p_name.paragraph_format.space_after = Pt(1)
    r_name = p_name.add_run(name)
    set_run_font(r_name, 18, bold=True)

    for contact in lines[1:3]:
        p = left.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(contact)
        set_run_font(run, 8.5)

    if PHOTO.exists():
        p_photo = right.paragraphs[0]
        p_photo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_photo.add_run().add_picture(str(PHOTO), width=Cm(2.35))

    return 3


def build_docx() -> None:
    """主流程。"""
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    for style_name in ("Normal", "Body Text", "List Paragraph"):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(7.6)

    lines = parse_markdown_lines()
    index = add_header(doc, lines)

    while index < len(lines):
        line = lines[index]
        if line.startswith("## "):
            add_section_heading(doc, line[3:])
        elif line.startswith("# "):
            pass
        else:
            is_entry_title = (
                " 20" in line
                and (
                    "吉林大学" in line
                    or "北京航空航天大学" in line
                    or "全国大学生电动方程式大赛" in line
                    or "奇绩创坛" in line
                    or "机器人与四足机器狗" in line
                    or "临舟科技有限公司" in line
                )
            )
            add_plain_paragraph(doc, line, size=7.5 if not is_entry_title else 7.9, bold=is_entry_title)
        index += 1

    doc.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    build_docx()
